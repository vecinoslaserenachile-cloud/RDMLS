import os
import glob
try:
    import fitz  # PyMuPDF
    from docx import Document
except ImportError:
    print("Por favor instala las dependencias: pip install PyMuPDF python-docx")
    exit(1)

TARGET_DIR = "C:/Users/estud/APP_LS_SEGURA/public/media/tano"
TEXTS_TO_REMOVE = ["UCEN", "Universidad Central", "NotebookLM", "notebooklm"]

def clean_docx(file_path):
    print(f"Procesando DOCX: {file_path}")
    try:
        doc = Document(file_path)
        modified = False
        
        # Replace in paragraphs
        for p in doc.paragraphs:
            for text in TEXTS_TO_REMOVE:
                if text.lower() in p.text.lower():
                    # Replace keeping case is hard in docx, we just do simple replace
                    # Actually, docx runs need to be replaced carefully
                    for run in p.runs:
                        if text.lower() in run.text.lower():
                            run.text = run.text.replace(text, "")
                            run.text = run.text.replace(text.upper(), "")
                            run.text = run.text.replace("Universidad Central", "")
                            run.text = run.text.replace("UCEN", "")
                            run.text = run.text.replace("NotebookLM", "")
                            modified = True
                            
        # Replace in tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        for run in p.runs:
                            for text in TEXTS_TO_REMOVE:
                                if text.lower() in run.text.lower():
                                    run.text = run.text.replace(text, "")
                                    run.text = run.text.replace(text.upper(), "")
                                    run.text = run.text.replace("Universidad Central", "")
                                    run.text = run.text.replace("UCEN", "")
                                    run.text = run.text.replace("NotebookLM", "")
                                    modified = True
        
        if modified:
            # Backup original
            backup_path = file_path + ".bak"
            if not os.path.exists(backup_path):
                os.rename(file_path, backup_path)
            doc.save(file_path)
            print(f"  -> Limpiado y guardado.")
        else:
            print(f"  -> Sin marcas encontradas.")
    except Exception as e:
        print(f"  -> Error procesando DOCX: {e}")

def clean_pdf(file_path):
    print(f"Procesando PDF: {file_path}")
    try:
        doc = fitz.open(file_path)
        modified = False
        
        for page in doc:
            # 1. Redact text
            for text in TEXTS_TO_REMOVE:
                # Search for text (case-insensitive)
                text_instances = page.search_for(text)
                for inst in text_instances:
                    page.add_redact_annot(inst, fill=(1, 1, 1)) # White fill
                    modified = True
            
            # Apply redactions
            page.apply_redactions()
            
            # 2. Draw white rectangles over common watermark/logo areas
            rect = page.rect
            # Bottom edge for NotebookLM watermark (bottom 5%)
            bottom_rect = fitz.Rect(0, rect.height - 30, rect.width, rect.height)
            page.draw_rect(bottom_rect, color=(1,1,1), fill=(1,1,1), overlay=True)
            
            # Top edge corners for UCEN logos
            # Top-left (width 20%, height 10%)
            top_left_rect = fitz.Rect(0, 0, rect.width * 0.2, rect.height * 0.1)
            page.draw_rect(top_left_rect, color=(1,1,1), fill=(1,1,1), overlay=True)
            
            # Top-right (width 20%, height 10%)
            top_right_rect = fitz.Rect(rect.width * 0.8, 0, rect.width, rect.height * 0.1)
            page.draw_rect(top_right_rect, color=(1,1,1), fill=(1,1,1), overlay=True)
            
            # We assume it modified something since we drew rectangles
            modified = True

        if modified:
            tmp_path = file_path + ".tmp"
            doc.save(tmp_path, garbage=4, deflate=True)
            doc.close()
            backup_path = file_path + ".bak"
            if not os.path.exists(backup_path):
                os.rename(file_path, backup_path)
            else:
                os.remove(file_path) # if backup exists, just overwrite original
            os.rename(tmp_path, file_path)
            print(f"  -> Limpiado y guardado.")
        else:
            doc.close()
            print(f"  -> Sin marcas encontradas.")
            
    except Exception as e:
        print(f"  -> Error procesando PDF: {e}")

if __name__ == "__main__":
    # Process DOCX
    for docx_file in glob.glob(os.path.join(TARGET_DIR, "*.docx")):
        clean_docx(docx_file)
        
    # Process PDF
    for pdf_file in glob.glob(os.path.join(TARGET_DIR, "*.pdf")):
        clean_pdf(pdf_file)
        
    print("Proceso finalizado.")
